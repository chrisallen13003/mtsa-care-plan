import streamlit as1: `app.py`**

```python
import streamlit as st
import json st
import json
import os
import os
from datetime
from datetime import datetime
from docx import datetime
from docx import Document
from docx import Document
from docx.shared import Pt
import.shared import Pt
import io
import anthrop io
import anthropicic

# Page

# Page config
st.set_page_ config
st.set_page_config(page_title="MTconfig(page_title="MTSA Care Plan Generator", page_icon="SA Care Plan Generator", page_icon="🏥", layout="wide")

#🏥", layout="wide")

# Title Title
st.title("
st.title("🏥 MTSA🏥 MTSA Anesthesia Care Plan Generator") Anesthesia Care Plan Generator")
st.markdown("###
st.markdown("### AI-Powered Custom AI-Powered Custom Care Plan Care Plan Creation Creation")")
st.markdown("---")

#
st.markdown("---")

# Sidebar for Sidebar for API API key key
with
with st.sidebar:
    st. st.sidebar:
    st.header("header("⚙️ Configuration⚙️ Configuration")
    api_key = st")
    api_key = st.text_input("Enter your Anthropic API Key.text_input("Enter your:", type="password", Anthropic API Key help="Get your:", type="password", help="Get your API key from console API key from console.anthropic.com")
    .anthropic.com")
    stst.markdown("---")
    st.markdown("---")
    st.markdown.markdown("###("###  📋📋 How How to Use:") to Use:")
    st.markdown("""
    1. Enter your API
    st.markdown("""
    1. Enter your API key ( key (sidebarsidebar)
    2. Type)
    2. Type a a brief brief patient patient scenario
    3. Click  scenario
    3. Click 'Generate Care Plan'
    4. Download your DOCX file'Generate Care Plan'
    4. Download your DOCX file
    """)
    """)
    st.markdown("---")
    st.markdown("---")
    st.markdown("### 
    st.markdown("### 💡 Example💡 Example Scenario:")
    st.code Scenario:")
    st.code("""48 yo F for lap chole with("""48 yo F for lap chole with HTN,  HTN, 
GERD,
GERD, hypothyroidism on  hypothyroidism on 
metoprolol, proton
metoprolol, protonix, synthroid.ix, synthroid. 
TIVA case.""") 
TIVA case.""")

# Main

# Main input input area
col area
col1, col2 = st.1, col2 = st.columns([columns([2, 1])

with col12, 1])

with col1:
    st.sub:
    st.subheader("📝 Patientheader("📝 Patient Scenario Input")
    patient_scenario Scenario Input")
    patient_scenario = st.text_area( = st.text_area(
        "Enter patient
        "Enter patient information information:",:",
        height
        height=150,
        placeholder="Example: 52=150,
        placeholder="Example yo M for TE: 52 yo M for TEEE an and cardiovd cardioversion, AFersion,ib, AFib, E EF 30%,F 30%, OS OSA, onA, on metoprolol an metoprolol and apd apixaban.ixaban. Deep Deep se sedation plannedation planned."
    )d."
    )

with col2:
    st.

with col2:
    st.subheader("subheader("📊 Optional📊 Optional Fields")
    student Fields")
    student_name = st.text_input_name = st.text_input("Student Name (("Student Name (optionaloptional):", "):", "")
    case")
    case_date_date = st.date_input("Case = st.date_input("Case Date:", Date:", datetime.now()) datetime.now())

# Generate button
if st.button

# Generate button
if st.button("("🎯 Generate Care Plan",🎯 Generate Care Plan", type type="primary",="primary", use_container_width=True):
    if not use_container_width=True):
    if not api_key:
        st. api_key:
        st.error("⚠️ Please enter yourerror("⚠️ Please enter your Anthropic API key Anthropic API key in the sidebar!") in the sidebar!")
    elif
    elif not patient_scenario. not patient_scenario.strip():
        st.error("⚠️ Please enter astrip():
        st.error("⚠️ Please enter a patient scenario!")
    else: patient scenario!")
    else:
        with
        with st.spinner(" st.spinner("🤖 AI🤖 AI is generating your custom is generating your custom care plan..."):
            try:
                # Initialize care plan..."):
            try:
                # Initialize Anthropic client
                client = anthrop Anthropic client
                client = anthropic.Anthropic(api_ic.Anthropic(api_key=api_key)
                key=api_key)
                
                # AI
                # AI Generation Generation Prompt
                prompt = f"""You Prompt
                prompt = f"""You are an expert anesthesia are an expert anesthesia provider provider creating a comprehensive an creating a comprehensive anesthesia care plan.esthesia care plan. Base Based on thed on the following following patient scenario, generate patient scenario, generate a complete, a complete, realistic care realistic care plan with plan with all all necessary details necessary details..

Patient Scenario: {patient_

Patient Scenario: {patient_scenario}

Generatescenario}

Generate a detaile a detailed JSONd JSON response response with the with the following structure. Fill following structure. Fill in realistic in realistic,, clin clinically appropriate valuesically appropriate values for ALL fields. for ALL fields. Use " Use "WWDDL" (L" (withinwithin define defined limits) for normald limits) for normal findings findings.. Be specific Be specific with with medications medications and doses and doses base based on the patient's weightd on the patient's weight and procedure and procedure.

Return ONLY vali.

Return ONLY valid JSON with thisd JSON with this exact exact structure: structure:

{{
  "age": 

{{
  "age": <number>,
  "gender": "<number>,
  "gender": "<M/F><M/F>",
  "height",
  "height__cm": <number>,
  "cm": <number>,
  "weight_kg": <number>,
  "bweight_kg": <number>,
  "bmi": <calculatemi": <calculated>,
  "id>,
  "idbwdb_w_kg": <calculated>,
  "kg": <calculated>,
  "asa": <1asa": <1-5-5>,
  "procedure_name": "<full>,
  "procedure_name": "<full procedure name>",
  "procedure procedure name>",
  "procedure_description": "<brief_description": "<brief highlights highlights:: positioning positioning, most, most stimulating stimulating parts parts, special, special considerations>",
  "n considerations>",
  "npopo__sincesince":": "<time description "<time description>",
  "npo_>",
  "npo_hourshours": <number": <number>,
  "allerg>,
  "allergies": "<listies": "<list or or N NKDA>",
  "routine_meds": "<list withKDA>",
  "routine indications>",
  "past_meds": "<list with indications>",
  "past_surgical_surgical_history": "<history_history": "<history or or None None>",
  "pre>",
  "preop_vs": {{op_vs": {{
    "bp":
    "bp": "<syst "<systolic/diastolic>olic/diastolic>",
    "hr": "<b",
    "hr": "<bpm>",
    "rrpm>",
    "rr": "<breaths/min>",": "<breaths/min>",
    "temp
    "temp": "<°C>",
    ": "<°C>",
    "s"spo2": "po2": "<%<%>> on RA"
  }},
  "pm on RA"
  }},
  "pmh": {{
    "cvh": {{
    "cv": "<cardiovascular history": "<cardiovascular history>",
    "resp>",
    "resp": "<respiratory history>",
    "c": "<respiratory history>",
    ns": "<neurological history>",
    "h"cnsep": "<hepatic history>",": "<neurological history>",
    "hep": "<hepatic history>",
    "gi
    "gi__gu": "<GI/Ggu": "<GI/GU history>",
    "U history>",
    "endo": "<endocrine history>endo": "<endocrine history>",
    "extrem",
    "extremities": "<musities": "<musculoskeletal>culoskeletal>",
    "other",
    "other": "<other relevant": "<other relevant history>"
  }},
   history>"
  }},
  "ros"ros": {{
    "cv": {{
    "cv": "<cardiovascular review": "<cardiovascular review>",
    "resp": ">",
    "resp": "<respiratory review>",
    "c<respiratory review>",
    "cns": "<neurological review>",ns": "<neurological review>",
    "hep": "<hepat
    "hep": "<hepatic review>",
    "giic review>",
    "gi_gu": "<GI/_gu": "<GI/GU review>",
    "GU review>",
    "extremextremities": "<extremities": "<extremity review>",
    "otherity review>",
    "other": "<other systems": "<other systems>"
  }},
  "cultural>"
  }},
  "cultural_religious_religious":": "<any "<any cultural or cultural or religious considerations or religious considerations or N N/A>",
  "labs/A>",
  "labs": {{
    "cbc":": {{
    "cbc": "<values or "<values or WDL or WDL or Not Not indicate indicated>",
    "chemd>",
    "chemistries": "<values or WDL or Not indicated>",
    "coistriesags": "<values or WDL or Not indicated>",
    "": "<values or WDL or Not indicated>",
    "coags": "<values or WDL or Not indicated>",
    "blooblood_typed_type": "<type or Not indicated>",
    "type": "<type_an or Not indicated>",
    "type_and_crossd_cross": "<units": "<units or Not indicated>",
    " or Not indicated>",
    "eekg": "<findingskg": "<findings or Not or Not indicated>",
    "echo": "<findings or Not indicated>", indicated>",
    "echo": "<findings or Not indicated>",
    "c
    "cxr": "<findings or Not indicated>xr": "<findings or Not indicated>",
    "h",
    "hccgg": "<result": "<result or N/A>",
    "other or N/A>",
    "other__studiesstudies": "<any": "<any other tests other tests or None or None>"
  }},
  "air>"
  }},
  "airway": {{
    "mallway": {{
    "mallampati": "<Iampati": "<I--IV>IV>",
    "thy",
    "thyromental_romental_distdist": "<distance in": "<distance in fingerbreadths>",
    " fingerbreadths>",
    "tmtmjj_problems": "<Yes_problems": "<Yes/No>",
    "missing/No>",
    "missing_loose_teeth": "<description_loose_teeth": "<description or or None>",
    "n None>",
    "npopo_since": "<time>_since": "<time>",
    "assessment",
    "assessment": "<overall": "<overall air airway assessment>"way assessment>"
  }},
  "an
  }},
  "anesthesia_plan": {{
    esthesia_plan": {{
    "type": "<MAC"type": "<MAC/GA/Regional/GA/Regional/combination/combination>",
    "air>",
    "airway_deviceway_device": "<ET": "<ETT/LMA/NaT/LMA/Nasal/Nonesal/None>",
    "regional>",
    "regional": "<technique": "<technique or or N N/A>",
    "pre/A>",
    "preop_mop_meds": "<medicationseds": "<medications with with doses>",
    "induction doses>",
    "induction_meds": "<medications with doses_meds": "<medications with doses base based on weight>",
    "maintenanced on weight>",
    "maintenance_meds": "<medications with doses_meds": "<medications with doses//infinfusion ratesusion rates>",
    "emergency>",
    "emergency_meds": "<rever_meds": "<reversal agents withsal agents with doses>",
    "me doses>",
    "med_considerations": "<drugd_considerations": "<drug interactions interactions,, patient patient--specific concernsspecific concerns>">"
  }},
  "flui
  }},
  "fluid_management": {{
    "ebd_management": {{
    "ebvv__mlml": ": <calculate<calculated>,
    "mad>,
    "mablbl_ml": <calculated>,
    _ml": <calculated>,
    "maintenance"maintenance_rate_rate": "": "<ml/hr>",
    "<ml/hr>",
    "nnpo_deficitpo_deficit": "<ml>",
    "": "<ml>",
    "insinsensible_lossensible_loss": "<ml/hr>",
    "thir": "<ml/d_space_hr>",
    "third_space_loss": "<ml/hr>",loss": "<ml/hr>",
    "fluid_plan
    "fluid_plan": "<crystall": "<crystalloid/colloid plan>"oid/colloid plan>"
  }},
  "case_considerations
  }},
  "case_considerations": [": [
    "
    "<consideration 1>",
    "<consideration<consideration 2>",
    "<consideration 1>",
    "<consideration 2>",
    "<consideration 3>",
    "<consideration 3>",
    "<consideration 4>",
    "<consideration 4>",
    "<consideration 5>",
    "<consideration 6>" 5>",
    "<consideration 6>"
  ],
  "intervent
  ],
  "interventions": [
    "<intervention/ions": [
    "<intervention/rationale 1>",
    "<intervention/rationale 2>",
    rationale 1>",
    "<intervention/rationale 2>",
    "<intervention/rationale 3>"<intervention/rationale 3>",
    "<intervention/rationale 4>",
    "<intervention",
    "<intervention/rationale 4>",
    "<intervention/rationale 5>",/rationale 5>",
    "<intervention/rationale 6>"
  ]
}}

Ensure all medication doses are weight-based and appropriate. Make
    "<intervention/rationale 6>"
  ]
}}

Ensure all medication doses are weight-based and appropriate. Make all values all values realistic and internally realistic and internally consistent.""" consistent."""

                # Call

                # Call Claude Claude API API
                message
                message = client.messages.create( = client.messages.create(
                    model="claude-opus
                    model="claude-opus--4-4-2020250250514514",
                    max_tokens=",
                    max_tokens=4000,
                    temperature=4000,
                    temperature=0.0.7,
                    messages7,
                    messages=[
                        {"role": "=[
                        {"role": "user", "content": prompt}user", "content": prompt}
                    ]
                )
                
                    ]
                )
                
                # Extract
                # Extract JSON JSON from from response
                response_text = message response
                response_text = message.content[0].text
                .content[0].text
                
                #
                # Try Try to parse JSON to parse JSON
                #
                # Sometimes Claude Sometimes Claude wr wraps JSONaps JSON in markdown in markdown code code blocks
                if blocks
                if " "```json" in response_text:```json" in response_text:
                    json
                    json_str = response_text.split_str = response_text.split("```json")[1].split("("```json")[1].split("```")[0].strip()
                elif```")[0].strip()
                elif "``` "```" in response_text:
                " in response_text:
                    json_str = response_text    json_str = response_text.split("```")[1].split.split("```")[1].split("```")[0].strip()("```")[0].strip()
                else:
                    json_
                else:
                    json_str = response_text.strip()str = response_text.strip()
                
                care
                
                care_plan__plan_datadata = json.loads(json_str = json.loads(json_str)
                
                # Display)
                
                # Display success success
                st.success("✅
                st.success("✅ Care plan Care plan generated successfully!")
                 generated successfully!")
                
                # Show preview
                # Show preview
                with
                with st.expander(" st.expander("📄 Preview Generate📄 Preview Generated Datad Data", expanded=True):
                    col",1, col2, col3 = st.columns( expanded=True):
                    col1, col2, col3 = st.columns(3)
                    with3)
                    with col1:
                        st.metric col1:
                        st.metric("Age("Age", care_plan_data.", care_plan_data.get("age", "N/Aget("age", "N/A"))
                        "))
                        st.metric("Gender", care_plan_data.get("gender", "N/A"))
                        st.metric("Gender", care_plan_data.get("gender", "N/A"))
                        st.metric("ASst.metric("ASA", care_plan_data.A", care_plan_data.get("asa", "N/get("asa", "N/A"))
                    with col2:
                        st.metric("A"))
                    with col2Height (cm)", care_plan_:
                        st.metric("Height (cm)", care_plan_data.get("height_cm",data.get("height_cm", "N/A"))
                         "N/A"))
                        st.metric("Weight (kg)",st.metric("Weight (kg)", care_plan_data.get(" care_plan_data.get("weight_kg", "N/Aweight_kg", "N/A"))
                        "))
                        st.metric("BMI", carest.metric("BMI", care_plan_data.get("bmi", "N/A"))_plan_data.get("bmi", "N/A"))
                    with col3:
                
                    with col3:
                        st.write        st.write("**("**ProcedureProcedure:**"):**")
                        st.write(care_plan_data.get("procedure
                        st.write(care_plan_data.get("procedure_name", "N/A"))_name", "N/A"))
                    
                    
                    st.markdown
                    st.markdown("---")
                    st.("---")
                    st.jsonjson(care_plan_data)(care_plan_data)
                
                
                # Create
                # Create Wor Word document
                std document
                st.sub.subheader("header("📥📥 Download Your Care Plan")
                
                # Fill Download Your Care Plan")
                
                # Fill template
                 template
                doc = filldoc = fill_template(care_plan_data_template(care_plan_data,, student_name, case_date) student_name, case_date)
                
                # Save
                
                # Save to to bytes bytes
                doc_io
                doc_io = io.BytesIO() = io.BytesIO()
                doc.save(doc_io
                doc.save(doc_io)
                doc_io.seek)
                doc_io.seek(0)
                
                #(0)
                
                # Download button Download button
                st.download_button(
                st.download_button(
                    label="📄
                    label="📄 Downloa Download DOCX",
                    data=d DOCX",
                    data=doc_io,
                    file_name=doc_io,
                    file_name=f"MTSA_Carf"MTSA_CarePlan_{careePlan_{care_plan_data.get('procedure_plan_data.get('procedure_name', 'case_name', 'case').').replace(' ', '_'replace(' ', '_')}_{datetime)}_{datetime.now().strftime('%Y%.now().strftime('%Y%m%d')m%d')}.docx",
                    mime}.docx",
                    mime="application/vnd.openxmlformats-office="application/vnd.openxmlformats-officedocument.wordprocessingml.document",document.wordprocessingml.document",
                    use_container_width
                    use_container_width=True
                )
                =True
                )
                
            except json
            .JSONDecodeError asexcept json.JSONDecodeError as e:
                st.error( e:
                st.error(f"f"❌ Error parsing AI❌ Error parsing AI response: {str(e)}")
                st.code(response response: {str(e)}")
                st.code(response_text)_text)
            
            except Exception as e:
                stexcept Exception as e:
                st.error(f"❌.error(f"❌ Error generating care plan: {str(e)}")
                st.exception Error generating care plan: {str(e)}")
                st.exception(e)(e)


def fill_template(data


def fill_template(data, student_name, case_date, student_name, case_date):
    """Fill the):
    """Fill the MTSA template with generate MTSA template with generated data"""
    
    # Create a new d data"""
    
    # Create a new document (document (you'll replaceyou'll replace this with template this with template loading)
    doc = Document() loading)
    doc = Document()
    
    # Add title
    
    
    # Add title
    title = doc.add_heading('MTSAtitle = doc.add_heading('MTSA Anesthesia Care Plan',  Anesthesia Care Plan', 0)
    0)
    
    # Header info
    # Header info
    header_table
    header_table = doc.add_table(rows=2 = doc.add_table(rows=2, cols=3, cols=3)
    header_table.cell)
    header_table.cell(0, 0).text =(0, 0).text = f f""STUDENTSTUDENT NAME: {student_name if NAME: {student_name if student_name else 'Click student_name else 'Click or or tap tap here to enter text.' here to enter text.'}"
    header_table.cell}"
    header_table.cell(0, 1).text = f"DATE(0, 1).text = f"DATE: {case_date.strftime: {case_date.strftime('%m('%m/%d/%Y')}"
    /%d/%Y')}"
    header_table.cell(0,header_table.cell(0, 2).text = f" 2).text = f"SURGICALSURGICAL PROCEDURE: {data.get('procedure_name', 'N/ PROCEDURE: {data.get('procedure_name', 'N/A')}"
    
    #A')}"
    
    # Patient Patient Data Data Section Section
    
    doc.add_heading('doc.add_heading('PATIENT DATA', PATIENT DATA', 1)
    patient1)
    patient_table = doc.add_table(rows=7_table = doc.add_table(rows=7, cols=2, cols=2)
    patient)
    patient_table.cell(0, _table.cell(0, 0).text = f"AG0).text = f"AGE: {data.get('ageE: {data.get('age', 'N/A')}"
    patient_table.cell(0, 1).text = f', 'N/A')}"
    patient_table.cell(0, 1).text = f"ALLERG"ALLERGIES: {data.get('allergIES: {data.get('allergies', 'N/A')ies', 'N/A')}"
    patient_table.cell}"
    patient_table.cell(1(1, 0).text = f", 0).text = f"GENDER: {data.get('GENDER: {data.get('gender', 'N/A')}"
    patient_table.cell(1, 1).text =gender', 'N/A')}"
    patient_table.cell(1, 1).text = f"ROUTINE f"ROUTINE MEDICATIONS: {data.get('routine_meds', 'N/A MEDICATIONS: {data.get('routine_meds', 'N/A')}"
    patient_table.')}"
    patient_table.cell(2, 0).textcell(2, 0).text = f"H = f"HTT: {data.get('height: {data.get('height_cm', 'N/A'_cm', 'N/A')} cm"
    patient_table)} cm"
    patient_table.cell(3.cell(3, 0).text = f", 0).text = f"WT: {data.get('weightWT: {data.get('weight_kg', 'N/A_kg', 'N/A')} kg"
    patient_')} kg"
    patient_table.cell(4, 0table.cell(4, 0).text = f"BMI:).text = f"BMI: {data.get('bmi', {data.get('bmi', 'N/A')}" 'N/A')}"
    patient_table.cell(5
    patient_table.cell(5, 0).text = f", 0).text = f"IDBW: {data.get('idbw_kg', 'N/A'IDBW: {data.get('idbw_kg', 'N/A')} kg"
    )} kg"
    
    # Vital
    # Vital Signs Signs
    doc.add_heading('PR
    doc.add_heading('PREANESTHEANESTHESIA VS', 1)
    vsESIA VS_data = data', 1)
    vs_data = data.get('preop_vs', {})
    vs_table =.get('preop_vs', {})
    vs_table = doc.add_table(rows= doc.add_table(rows=5, cols=2)
    vs_table.cell(0, 5, cols=2)
    vs_table.cell(0, 0).text = f"BP: {vs_data.get('bp0).text = f"BP: {vs_data.get('bp', 'N/A')}"', 'N/A')}"
    vs_table.cell(1
    vs_table.cell(1, 0).text = f", 0).text = f"HR: {vs_data.get('hr', 'N/AHR: {vs_data.get('hr', 'N/A')}"
    vs_table.cell(2, 0).text = f"RR: {vs_data.get('rr', 'N/A')}"')}"
    vs_table.cell(2, 0).text = f"RR: {vs_data.get('rr', 'N/A')}"
    vs_table.cell(3, 0).text = f"
    vs_table.cell(3, 0).text = f"TEMP: {vs_data.get('temp', 'N/TEMP: {vs_data.get('temp', 'N/A')}"
    vs_tableA')}"
    vs_table.cell(4, 0)..cell(4, 0).text = f"SAO2: {vs_data.get('spo2', 'N/A')}"
    text = f"SAO2: {vs_data.get('spo2', '
    # AirN/A')}"
    
    # Airway Assessment
    doc.add_heading('AIRWAY ASSESSMENT', 1)
    airway_way Assessment
    doc.add_data = data.get('airway', {})
    doc.add_paragraph(f"heading('AIRWAY ASSESSMENT', 1)
    airway_data = data.get('airway', {})
    doc.add_paragraph(f"MALLAMPATI: {airway_MALLAMPATI: {airway_data.get('mallampati',data.get('mallampati', 'N/A')}") 'N/A')}")
    doc.add_paragraph(f
    doc.add_paragraph(f"THYROMENTAL DIST: {airway_data.get('thyromental_dist', "THYROMENTAL DIST: {airway_data.get('thyromental_dist', 'N/A')}")
    doc.add_paragraph(f"TMJ PROBLEMS: {airway'N/A')}")
    doc.add_paragraph(f"_data.get('tmj_problems', 'N/A')TMJ PROBLEMS: {airway_data.get('tmj_}")
    doc.add_paragraph(f"MISSING/problems', 'N/A')}")
    doc.add_paragraphLOOSE TEETH: {airway_data.get('missing_loose(f"MISSING/LOOSE TEETH: {airway_data.get('missing_loose_teeth', 'N/A')}")
    doc.add_paragraph(f"NPO SINCE_teeth', 'N/A')}")
    doc.add_paragraph(f"NPO SINCE: {airway_data.get('npo_since', 'N/A')}")
    : {airway_data.get('npo_since', 'N/A')}")
    doc.add_paragraph(f"AIRWAY:doc.add_paragraph(f"AIRWAY: {airway_data.get('assessment', 'N/A') {airway_data.get('assessment', 'N/A')}")
    
    # Past Medical History
    doc.add_heading('PAST MEDICAL HISTORY}")
    
    # Past Medical History
    doc.add_heading('PAST MEDICAL HISTORY', 1)
    pm', 1)
    pmh_data = data.get('pmh', {})
    doc.add_paragraph(f"CVh_data = data.get('pmh', {})
    : {pmh_data.get('cv', 'WDL')}")
    doc.add_paragraph(f"RESP: {pmh_data.getdoc.add_paragraph(f"CV: {pmh_data.get('cv', 'WDL')}")
    doc.add_paragraph(f"RESP: {pmh_data.get('resp', 'WDL')}")
    doc.add_('resp', 'WDL')}")
    doc.add_paragraph(f"CNS: {pmh_data.paragraph(f"CNS: {pmh_data.get('cns', 'Wget('cns', 'WDL')}")
    doc.add_paragraph(f"HEP: {pmh_data.get('hep', 'WDL')}")
    DL')}")
    doc.add_paragraph(f"HEP: {pmh_data.get('hep', 'WDL')}")
    doc.add_paragraph(f"GI/GU: {pmh_data.get('gi_gu', 'WDL')}")
    doc.add_paragraphdoc.add_paragraph(f"GI/GU: {pmh_data.get('gi_gu', 'WDL')}")
    doc.add_paragraph(f"EXTREM: {pmh_data.get('extremities', 'WDL')}")
    doc.add_paragraph(f"OTHER: {pmh_data.get('other', 'WDL')}")
    
    # Review of Systems
    doc.add_heading('REVIEW OF SYSTEMS', 1)
    ros_data = data.get('ros', {(f"EXTREM: {pmh_data.get('extremities', 'WDL')}")
    doc.add_paragraph(f"OTHER: {pmh_data.get('other', 'WDL')}")
    
    # Review of Systems
    doc.add_heading('REVIEW OF SYSTEMS', 1)
    ros_data = data.get('ros', {})
    doc.add_paragraph(f"CV: {ros_data.get('cv', 'WDL')}")
    doc.add_paragraph(f"RESP: {ros_data.get('resp', 'WDL')}")
    doc.add_paragraph(f"CNS: {ros_data.get('cns', 'WDL')}")
    doc.add_paragraph(f"HEP: {ros_data.get('hep', 'WDL')}")
    doc.add_paragraph(f"GI/GU: {ros_data.get('gi_gu', 'WDL')}")
    doc.add_paragraph(f"EXTREMITIES: {ros_data.get('extremities', 'WDL')}")
    doc.add_paragraph(f"OTHER: {ros_data.get('other', 'WDL')}")})
    doc.add_paragraph(f"CV: {ros_data.get('cv', 'WDL')}")
    doc.add_paragraph(f"RESP: {ros_data.get('resp', 'WDL')}")
    doc.add_paragraph(f"CNS: {ros_data.get('cns', 'WDL')}")
    doc.add_paragraph(f"HEP: {ros_data.get('hep', 'WDL')}")
    doc.add_paragraph(f"GI/GU: {ros_data.get('gi_gu', 'WDL')}")
    doc.add_paragraph(f"EXTREMITIES: {ros_data.get('extremities', 'WDL')}")
    doc.add_paragraph(f"OTHER: {ros_data.get('other', 'WDL')}")
    
    # Labs
    doc.add_heading('DIAGNOSTIC & LAB STUDIES', 1)
    labs_data = data.get('labs', {})
    doc.add_paragraph(f"CBC: {labs_data.get('cbc', 'Not indicated')}")
    doc.add_paragraph(f"CHEMISTRIES: {labs_data.get('chemistries', 'Not indicated')}")
    doc.add_paragraph(f"COAGS: {labs_data.get('coags', 'Not indicated')}")
    doc.add_paragraph(f"BLOOD TYPE: {labs_data.get('blood_type', 'Not indicated')}")
    doc.add_paragraph(f"T&C: {labs_data.get('type_and_cross', 'Not indicated')}")
    doc.add_paragraph(f"EKG: {labs_data.get('ekg', 'Not indicated')}")
    doc.add_paragraph(f"ECHO: {labs_data.get('echo', 'Not indicated')}")
    doc.add_paragraph(f"CXR: {labs_data.get('cxr', 'Not indicated')}")
    doc.add_paragraph(f"HCG: {labs_data.get('hcg', 'N/A')}")
    
    # Anesthesia Plan
    
    # Labs
    doc.add_heading('DIAGNOSTIC & LAB STUDIES', 1)
    labs_data = data.get('labs', {})
    doc.add_paragraph(f"CBC: {labs_data.get('cbc', 'Not indicated')}")
    doc.add_paragraph(f"CHEMISTRIES: {labs_data.get('chemistries', 'Not indicated')}")
    doc.add_paragraph(f"COAGS: {labs_data.get('coags', 'Not indicated')}")
    doc.add_paragraph(f"BLOOD TYPE: {labs_data.get('blood_type', 'Not indicated')}")
    doc.add_paragraph(f"T&C: {labs_data.get('type_and_cross', 'Not indicated')}")
    doc.add_paragraph(f"EKG: {labs_data.get('ekg', 'Not indicated')}")
    doc.add_paragraph(f"ECHO: {labs_data.get('echo', 'Not indicated')}")
    doc.add_paragraph(f"CXR: {labs_data.get('cxr', 'Not indicated')}")
    doc.add_paragraph(f"HCG: {labs_data.get('hcg', 'N/A')}")
    
    # Anesthesia Plan
    doc.add_heading('ANESTHESIA PLAN', 1)
    anesthesia_data = data.get('anesthesia_plan', {})
    doc.add_paragraph(f"TYPE OF ANESTHETIC: {anesthesia_data.get('type', 'N/A')}")
    doc.add_paragraph(f"AIRWAY DEVICE: {anesthesia_data.get('airway_device', 'N/A')}")
    doc.add_paragraph(f"REGIONAL: {anesthesia_data.get('regional', 'N/A')}")
    doc.add_paragraph(f"\nPREOPERATIVE MEDS & DOSES:\n{anesthesia_data.get('preop_meds', 'N/A')}")
    doc.add_paragraph(f"\nINDUCTION MEDS & DOSES:\n{anesthesia_data.get('induction_meds', 'N/A')}")
    doc.add_paragraph(f"\nMAINTENANCE MEDS & DOSES:\n{anesthesia_data.get('maintenance_meds', 'N/A')}")
    doc.add_paragraph(f"\nEMERGENCY PLAN/REVERSAL MEDS & DOSES:\n{anesthesia_data.get('emergency_meds', 'N/A')}")
    doc.add_paragraph(f"\nMEDICATION CONSIDERATIONS:\n{anesthesia_data.get('med_considerations', 'N/A')}")
    
    # Procedure
    doc.add_heading('PROCEDURE', 1)
    doc.add_paragraph(data.get('procedure_description', 'N/A'))
    
    # Case Considerations
    doc.add_heading('POTENTIAL PATIENT-SPECIFIC OR CASE-SPECIFIC CONSIDERATIONS', 1)
    considerations = data.get('case_considerations', [])
    for i, consideration in enumerate(considerations, 1):
        doc.add_paragraph(f"{i}. {consideration}")
    
    # Interventions
    doc.add_heading('INTERVENTIONS & RATIONALE', 1)
    interventions = data.get('interventions', [])
    for i, intervention in enumerate(interventions, 1):
        doc.add_paragraph(f"{i}. {intervention}")
    
    # Fluid Management
    doc.add_heading('BLOOD AND BODY FLUID REQUIREMENTS', 1)
    fluid_data = data.get('fluid_management', {})
    doc.add_paragraph(f"Estimated blood volume (EBV): {fluid_data.get('ebv_ml', 'N/A')} mL")
    doc.add_paragraph(f"Maximum
    doc.add_heading('ANESTHESIA PLAN', 1)
    anesthesia_data = data.get('anesthesia_plan', {})
    doc.add_paragraph(f"TYPE OF ANESTHETIC: {anesthesia_data.get('type', 'N/A')}")
    doc.add_paragraph(f"AIRWAY DEVICE: {anesthesia_data.get('airway_device', 'N/A')}")
    doc.add_paragraph(f"REGIONAL: {anesthesia_data.get('regional', 'N/A')}")
    doc.add_paragraph(f"\nPREOPERATIVE MEDS & DOSES:\n{anesthesia_data.get('preop_meds', 'N/A')}")
    doc.add_paragraph(f"\nINDUCTION MEDS & DOSES:\n{anesthesia_data.get('induction_meds', 'N/A')}")
    doc.add_paragraph(f"\nMAINTENANCE MEDS & DOSES:\n{anesthesia_data.get('maintenance_meds', 'N/A')}")
    doc.add_paragraph(f"\nEMERGENCY PLAN/REVERSAL MEDS & DOSES:\n{anesthesia_data.get('emergency_meds', 'N/A')}")
    doc.add_paragraph(f"\nMEDICATION CONSIDERATIONS:\n{anesthesia_data.get('med_considerations', 'N/A')}")
    
    # Procedure
    doc.add_heading('PROCEDURE', 1)
    doc.add_paragraph(data.get('procedure_description', 'N/A'))
    
    # Case Considerations
    doc.add_heading('POTENTIAL PATIENT-SPECIFIC OR CASE-SPECIFIC CONSIDERATIONS', 1)
    considerations = data.get('case_considerations', [])
    for i, consideration in enumerate(considerations, 1):
        doc.add_paragraph(f"{i}. {consideration}")
    
    # Interventions
    doc.add_heading('INTERVENTIONS & RATIONALE', 1)
    interventions = data.get('interventions', [])
    for i, intervention in enumerate(interventions, 1):
        doc.add_paragraph(f"{i}. {intervention}")
    
    # Fluid Management
    doc.add_heading('BLOOD AND BODY FLUID REQUIREMENTS', 1)
    fluid_data = data.get('fluid_management', {})
    doc.add_paragraph(f"Estimated blood volume (EBV): {fluid_data.get('ebv_ml', 'N/A')} mL")
    doc.add_paragraph(f"Maximum Allowable blood loss (MABL): {fluid_data.get('mabl_ml', 'N/A')} mL")
    doc.add_paragraph(f"Maintenance IVF rate: {fluid_data.get('maintenance_rate', 'N/A')}")
    doc.add_paragraph(f"NPO Deficit: {fluid_data.get('npo_deficit', 'N/A')}")
    doc.add_paragraph(f"Insensible Loss: {fluid_data.get('insensible_loss', 'N/A')}")
    doc.add_paragraph(f"3rd Space Loss: {fluid_data.get('third_space_loss', 'N/A')}")
    doc.add_paragraph(f"Fluid Plan: {fluid_data.get('fluid_plan', 'N/A')}")
    
    return doc


# Footer
st.markdown("---")
st.markdown("### 🔒 Privacy Notice")
st.caption("This app generates educational simulations only. No patient data is stored. All content is created transiently and deleted after download.")
